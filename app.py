import streamlit as st
import fitz  # PyMuPDF
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.embeddings import FastEmbedEmbeddings
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain_groq import ChatGroq
import os
import matplotlib.pyplot as plt
from docx import Document  # For handling .docx files
from PIL import Image  # For handling images

# Function to extract text from a PDF file
def load_pdf_text(pdf_path):
    doc = fitz.open(pdf_path)
    text_data = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text_data.append(page.get_text("text"))

    return text_data

# Function to extract text from a .docx file
def load_docx_text(docx_path):
    doc = Document(docx_path)
    text_data = []

    for para in doc.paragraphs:
        text_data.append(para.text)

    return "\n".join(text_data)

# Function to extract keywords from text
def extract_keywords(text):
    # Placeholder for actual NLP logic to extract keywords
    # For demonstration, we split the text into words
    return set(text.lower().split())

# Function to compare resume and job description
def compare_resume_job_description(resume_text, job_description):
    # Extract keywords from resume and job description
    resume_keywords = extract_keywords(resume_text)
    job_keywords = extract_keywords(job_description)

    # Compare keywords
    matched_keywords = resume_keywords.intersection(job_keywords)
    mismatched_keywords = job_keywords.difference(resume_keywords)

    return matched_keywords, mismatched_keywords

# Function to calculate ATS score
def calculate_ats_score(matched_keywords, total_keywords):
    if total_keywords == 0:
        return 0
    return (len(matched_keywords) / total_keywords) * 100

# Function to generate suggestions based on mismatched keywords
def generate_suggestions(mismatched_keywords):
    suggestions = []
    if mismatched_keywords:
        # Limit the number of mismatched keywords to 7
        max_keywords_to_show = 7
        if len(mismatched_keywords) > max_keywords_to_show:
            suggestions.append(f"Add the following keywords to your resume (showing {max_keywords_to_show} of {len(mismatched_keywords)}): " + ", ".join(list(mismatched_keywords)[:max_keywords_to_show]))
        else:
            suggestions.append("Add the following keywords to your resume: " + ", ".join(mismatched_keywords))
    return suggestions

# Function to edit the resume and add new keywords
def edit_resume(docx_path, new_keywords, output_path):
    # Load the .docx file
    doc = Document(docx_path)

    # Find and update the skills section
    for para in doc.paragraphs:
        if "skills" in para.text.lower():  # Identify the skills section
            # Remove double commas and colons from the existing skills
            existing_skills = para.text.replace(",,", ",").replace(":", "").strip()
            
            # Append new keywords to the existing skills
            updated_skills = existing_skills + ", " + ", ".join(new_keywords)
            
            # Remove any double commas and colons from the updated skills
            updated_skills = updated_skills.replace(",,", ",").replace(":", "").strip()
            
            # Update the paragraph text
            para.text = updated_skills

    # Save the updated .docx file
    doc.save(output_path)

# Function to convert .docx to PDF while preserving page structure
def convert_docx_to_pdf(docx_path, output_path):
    # Use PyMuPDF to convert .docx to PDF
    doc = fitz.open()
    page = doc.new_page()

    # Load text from .docx
    text = load_docx_text(docx_path)

    # Insert text into the PDF
    page.insert_text((50, 50), text, fontsize=12, fontname="helv", color=(0, 0, 0))

    # Save the PDF
    doc.save(output_path)
    doc.close()

# Function to extract images from PDF pages
def extract_images_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    images = []

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pix = page.get_pixmap()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        images.append(img)

    return images

# Streamlit UI setup
st.set_page_config(page_title="AI-Powered Tailored Resume Generator", layout="wide")

# Custom CSS for professional UI and background image
st.markdown(
    """
    <style>
    .main {
        background-image: url('https://raw.githubusercontent.com/your_username/your_repo/main/static/v.jpg');
        background-size: cover;
        background-position: center;
        background-repeat: no-repeat;
        padding: 2rem;
        color: white;  /* Ensure text is readable on the background */
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# User authentication
def authenticate(username, password):
    # Placeholder for actual authentication logic
    # Replace with a database or secure authentication mechanism
    users = {"admin": "password123", "user": "resume123"}
    return users.get(username) == password

# Login page
def login_page():
    st.title("Login to Resume Generator")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    if st.button("Login"):
        if authenticate(username, password):
            st.session_state["logged_in"] = True
            st.session_state["username"] = username
            st.success("Logged in successfully!")
        else:
            st.error("Invalid username or password")

# Logout functionality
def logout():
    st.session_state["logged_in"] = False
    st.session_state["username"] = None
    st.success("Logged out successfully!")

# Main application
def main_app():
    st.title("AI-Powered Tailored Resume Generator📄")
    st.write(f"Welcome, {st.session_state['username']}!")

    uploaded_file = st.file_uploader("Upload your resume (PDF or DOCX)", type=["pdf", "docx"])

    if uploaded_file is not None:
        with st.spinner("Processing your resume..."):
            file_path = "uploaded_resume." + uploaded_file.name.split(".")[-1]
            with open(file_path, "wb") as f:
                f.write(uploaded_file.read())

            # Load text from the uploaded file
            if uploaded_file.name.endswith(".pdf"):
                texts = load_pdf_text(file_path)
                # Extract images from PDF
                original_images = extract_images_from_pdf(file_path)
            elif uploaded_file.name.endswith(".docx"):
                texts = load_docx_text(file_path)
                # Convert DOCX to images
                original_images = extract_images_from_pdf(file_path)  # Placeholder for DOCX to images
            resume_text = " ".join(texts) if isinstance(texts, list) else texts

            # Prepare text chunks for embedding
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
            text_chunks = text_splitter.create_documents([resume_text])
            text_embeddings = FastEmbedEmbeddings(model_name="BAAI/bge-small-en-v1.5")
            text_db = FAISS.from_documents(text_chunks, text_embeddings)

            # Setup language model for Q&A
            os.environ['GROQ_API_KEY'] = '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'
            llm = ChatGroq(model_name='llama-3.1-8b-instant')
            memory = ConversationBufferMemory(memory_key='chat_history', return_messages=False)
            retriever = text_db.as_retriever(search_type="similarity", search_kwargs={"k": 3})
            qa = ConversationalRetrievalChain.from_llm(
                llm=llm,
                memory=memory,
                retriever=retriever
            )

            st.success("Resume processed successfully!")

            # Sidebar for extracted text
            st.sidebar.subheader("📜 Extracted Text")
            with st.sidebar.expander("View Extracted Text"):
                st.write(resume_text)

            # Tabs for asking questions and exploring text
            tab1, tab2 = st.tabs(["💬 Ask Questions/Compare with Job Description", "📜 Explore Text"])

            with tab1:
                query = st.text_input("Give the job description:")
                if st.button("Submit"):
                    if query:
                        with st.spinner("Processing your input..."):
                            # Use the conversational retrieval chain to process the query
                            result = qa({"question": query})
                            answer = result['answer']

                            # Create a container for the generated text
                            with st.container():
                                st.write("### 💡 Answer:")
                                st.write(answer)

                                # Compare resume and job description
                                matched_keywords, mismatched_keywords = compare_resume_job_description(resume_text, query)

                                # Calculate ATS score for the old resume
                                total_keywords = len(matched_keywords) + len(mismatched_keywords)
                                old_ats_score = calculate_ats_score(matched_keywords, total_keywords)

                                st.subheader("Comparison Results")
                                col1, col2 = st.columns(2)
                                with col1:
                                    st.write("### Matched Keywords")
                                    with st.expander(f"View {len(matched_keywords)} matched keywords"):
                                        st.write(", ".join(matched_keywords))
                                with col2:
                                    st.write("### Mismatched Keywords")
                                    with st.expander(f"View {len(mismatched_keywords)} mismatched keywords"):
                                        st.write(", ".join(mismatched_keywords))

                                # Pie chart for detailed keyword-wise visualization
                                st.subheader("Keyword-wise Similarity Breakdown")
                                labels = ['Matched Keywords', 'Mismatched Keywords']
                                sizes = [len(matched_keywords), len(mismatched_keywords)]
                                colors = ['#4CAF50', '#FF4500']

                                # Plot the pie chart
                                fig1, ax1 = plt.subplots()
                                ax1.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90)
                                ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
                                st.pyplot(fig1)

                                # Overall similarity graph
                                st.subheader("Overall Similarity")
                                if total_keywords > 0:
                                    overall_labels = ['Similarity', 'Mismatch']
                                    overall_sizes = [len(matched_keywords), len(mismatched_keywords)]
                                    overall_colors = ['#4CAF50', '#FF4500']

                                    fig2, ax2 = plt.subplots()
                                    ax2.pie(overall_sizes, labels=overall_labels, colors=overall_colors, autopct='%1.1f%%', startangle=90)
                                    ax2.axis('equal')
                                    st.pyplot(fig2)
                                else:
                                    st.write("No data available for overall similarity.")

                                # Generate and display suggestions
                                st.subheader("Suggestions to Improve Your Resume")
                                suggestions = generate_suggestions(mismatched_keywords)
                                if suggestions:
                                    for suggestion in suggestions:
                                        st.write(f"- {suggestion}")
                                else:
                                    st.write("Your resume is well-aligned with the job description! 🎉")

                                # Edit the resume and create a new DOCX
                                if mismatched_keywords:
                                    updated_docx_path = "updated_resume.docx"
                                    edit_resume(file_path, mismatched_keywords, updated_docx_path)
                                    st.success("Updated resume DOCX created successfully!")

                                    # Convert the updated .docx to PDF
                                    output_pdf_path = "updated_resume.pdf"
                                    convert_docx_to_pdf(updated_docx_path, output_pdf_path)

                                    # Calculate ATS score for the updated resume
                                    updated_resume_text = load_docx_text(updated_docx_path)
                                    updated_matched_keywords, _ = compare_resume_job_description(updated_resume_text, query)
                                    updated_ats_score = calculate_ats_score(updated_matched_keywords, total_keywords)

                                    # Display ATS scores
                                    st.subheader("ATS Scores")
                                    st.write(f"- Resume ATS Score: {updated_ats_score:.2f}%")

                                    # Provide download links
                                    with open(updated_docx_path, "rb") as f:
                                        st.download_button(
                                            label="Download Updated Resume (DOCX)",
                                            data=f,
                                            file_name="updated_resume.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                        )

                                    # Extract images from the original and updated PDFs
                                    original_images = extract_images_from_pdf(file_path)
                                    updated_images = extract_images_from_pdf(output_pdf_path)

                                    # Display side-by-side preview of original and updated pages
                                    st.subheader("Preview of Original and Updated Resume Pages")
                                    for i, (original_img, updated_img) in enumerate(zip(original_images, updated_images)):
                                        col1, col2 = st.columns(2)
                                        with col1:
                                            st.image(original_img, caption=f"Original Page {i + 1}", width=300)  # Fixed width
                                        with col2:
                                            st.image(updated_img, caption=f"Updated Page {i + 1}", width=300)  # Fixed width

            with tab2:
                st.subheader("📜 Extracted Text")
                st.write(resume_text)

    else:
        st.info("Upload a PDF or DOCX file to get started!")

# Main logic
if "logged_in" not in st.session_state:
    st.session_state["logged_in"] = False

if st.session_state["logged_in"]:
    if st.sidebar.button("Logout"):
        logout()
    main_app()
else:
    login_page()